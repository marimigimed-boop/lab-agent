"""
word_builder — единый шаблон Word для всех агентов интерпретации анализов.

Использование:
    from src.utils.word_builder import build_lab_word_bytes, build_email_with_word

    docx_bytes = build_lab_word_bytes(
        patient_name="Иванов А Б",
        date_str="03.05.2026",
        analysis_count=3,
        interpretation_text="... markdown от Claude ...",
    )

    raw_email = build_email_with_word(
        from_addr="moimed23@mail.ru",
        to_addr="marimigi@mail.ru",
        subject="Интерпретация — Иванов А Б",
        patient_name="Иванов А Б",
        date_str="03.05.2026",
        analysis_count=3,
        interpretation_text="... markdown от Claude ...",
    )
"""

import email as email_lib
import email.utils
import re
from datetime import datetime
from email.message import EmailMessage
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

# ── Константы шаблона ─────────────────────────────────────────────────────────
CLINIC_TITLE = "ЦЕНТР СЕМЕЙНОЙ МЕДИЦИНЫ И РЕАБИЛИТАЦИИ"
MOIMED_LINE  = (
    "«МОЙ МЕД»  МО, г. Одинцово, Красногорское шоссе, 10."
    "  Тел: 8 800 700 60 75 / 8 499 647 42 16"
    "  moimed23@mail.ru  moimed.com"
)
LITEMED_LINE = (
    "«ЛАЙТМЕД»  МО, г. Мытищи, Осташковское шоссе, вл. 22, стр. 1/1, эт. 2"
    "  Тел: 8 (800) 700-60-75 / 8 (499) 647-56-89"
    "  litemed@mail.ru  lite-med.ru"
)
SEPARATOR = "─" * 80
DOCTOR    = "Мигинеишвили Мария Давидовна"

BLUE  = RGBColor(0x1a, 0x6a, 0xc8)
GREEN = RGBColor(0x5c, 0xa8, 0x00)


# ── Вспомогательные функции ──────────────────────────────────────────────────
def _para(doc: Document, text: str = "", *, bold=False, size_pt: float = 11,
          color: RGBColor | None = None, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return p


def _separator(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SEPARATOR)
    run.font.size = Pt(7)


# ── Заголовок документа (шапка клиники) ──────────────────────────────────────
def _add_header(doc: Document, patient_name: str, date_str: str,
                analysis_count: int, doctor: str = DOCTOR) -> None:
    # Строка 1 — название клиники
    _para(doc, CLINIC_TITLE, bold=True, size_pt=11, color=BLUE,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    # Строка 2 — МОЙ МЕД
    _para(doc, MOIMED_LINE, size_pt=8, color=BLUE,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    # Строка 3 — ЛАЙТМЕД
    _para(doc, LITEMED_LINE, size_pt=8, color=GREEN,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    # Разделитель
    _separator(doc)

    # Строка 4 — данные пациента
    info = (
        f"Пациент: {patient_name}"
        f"   |   Дата: {date_str}"
        f"   |   Анализов: {analysis_count}"
        f"   |   Врач: {doctor}"
    )
    _para(doc, info, bold=True, size_pt=9, color=BLUE,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    # Разделитель
    _separator(doc)

    # Пустая строка
    doc.add_paragraph()


# ── Разбор Markdown → Word ────────────────────────────────────────────────────
def _is_table_header_sep(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return all(re.fullmatch(r"[-: ]+", c) for c in cells if c)


def _render_markdown(doc: Document, text: str) -> None:
    """Вставляет markdown-текст от Claude в документ Word."""
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # H2 / H3
        if line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = BLUE
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(11)
            i += 1
            continue

        # Таблица Markdown
        if line.startswith("| ") and line.endswith(" |"):
            # Собираем все строки таблицы
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                if not _is_table_header_sep(lines[i]):
                    table_lines.append(lines[i])
                i += 1

            if not table_lines:
                continue

            # Парсим ячейки
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                rows.append(cells)

            if rows:
                col_count = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=col_count)
                table.style = "Table Grid"
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < col_count:
                            cell = table.cell(ri, ci)
                            cell.text = cell_text
                            # Первая строка (заголовок) — жирная
                            if ri == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                doc.add_paragraph()
            continue

        # Горизонтальный разделитель
        if re.fullmatch(r"[-─═]+", line.strip()):
            _separator(doc)
            i += 1
            continue

        # Жирный текст **text**
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            i += 1
            continue

        # Обычная строка (возможно с inline **bold**)
        if line.strip():
            p = doc.add_paragraph()
            # Inline bold: разбиваем на части
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph()

        i += 1


# ── Публичный API ─────────────────────────────────────────────────────────────
def build_lab_word_bytes(
    patient_name: str,
    date_str: str,
    analysis_count: int,
    interpretation_text: str,
    doctor: str = DOCTOR,
) -> bytes:
    """
    Создаёт Word-документ по шаблону клиники.
    Возвращает байты .docx (для сохранения в файл или вложение).
    """
    doc = Document()

    # Поля страницы (A4, узкие)
    section = doc.sections[0]
    section.top_margin    = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin   = Pt(54)
    section.right_margin  = Pt(54)

    _add_header(doc, patient_name, date_str, analysis_count, doctor)
    _render_markdown(doc, interpretation_text)

    # Подпись в конце
    doc.add_paragraph()
    _separator(doc)
    _para(
        doc,
        f"Документ сформирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        f"  |  Врач: {doctor}",
        size_pt=8,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_multi_patient_word_bytes(
    patients: list[dict],
    title: str = "ИНТЕРПРЕТАЦИЯ ЛАБОРАТОРНЫХ АНАЛИЗОВ",
    doctor: str = DOCTOR,
) -> bytes:
    """
    Создаёт Word-документ для нескольких пациентов.
    Каждый пациент — отдельная страница с шапкой клиники.

    patients: список словарей с ключами:
        name           - ФИО пациента
        date           - дата анализов
        order          - номер заявки (опционально)
        analysis_count - количество анализов
        interpretation  - текст интерпретации (markdown)
    """
    doc = Document()

    # Поля страницы
    section = doc.sections[0]
    section.top_margin    = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin   = Pt(54)
    section.right_margin  = Pt(54)

    # Титульная страница — оглавление
    _para(doc, title, bold=True, size_pt=14, color=BLUE,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _para(doc, f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
          size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"Врач: {doctor}", size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    _separator(doc)
    _para(doc, "ПАЦИЕНТЫ В ДАННОМ ДОКУМЕНТЕ:", bold=True, size_pt=10, color=BLUE)
    for i, p in enumerate(patients, 1):
        order_info = f"  (заявка {p['order']})" if p.get("order") else ""
        doc.add_paragraph(f"{i}. {p['name']}{order_info}   —   {p['date']}")

    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    # Page break
    p_elem = doc.add_paragraph()
    p_elem._p.append(_OE("w:pageBreak"))

    for idx, patient in enumerate(patients):
        _add_header(
            doc,
            patient_name=patient["name"],
            date_str=patient["date"],
            analysis_count=patient.get("analysis_count", 1),
            doctor=doctor,
        )
        # Подзаголовок с номером заявки
        if patient.get("order"):
            _para(doc, f"Заявка №{patient['order']}", size_pt=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_paragraph()

        _render_markdown(doc, patient["interpretation"])

        # Разрыв страницы между пациентами (кроме последнего)
        if idx < len(patients) - 1:
            p_elem = doc.add_paragraph()
            p_elem._p.append(_OE("w:pageBreak"))

    # Подпись в конце
    doc.add_paragraph()
    _separator(doc)
    _para(
        doc,
        f"Документ сформирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        f"  |  Врач: {doctor}",
        size_pt=8,
        color=BLUE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_email_with_word(
    from_addr: str,
    to_addr: str,
    subject: str,
    patient_name: str,
    date_str: str,
    analysis_count: int,
    interpretation_text: str,
    doctor: str = DOCTOR,
) -> bytes:
    """
    Строит RFC-822 письмо с Word-вложением (.docx по шаблону клиники).
    Возвращает байты для IMAP append.
    """
    docx_bytes = build_lab_word_bytes(
        patient_name=patient_name,
        date_str=date_str,
        analysis_count=analysis_count,
        interpretation_text=interpretation_text,
        doctor=doctor,
    )

    safe_name = re.sub(r"[^\w\- ]", "", patient_name)[:50].strip()
    filename  = f"Анализы_{safe_name}_{date_str}.docx".replace(" ", "_")

    msg = EmailMessage()
    msg["From"]    = from_addr
    msg["To"]      = to_addr
    msg["Subject"] = subject
    msg["Date"]    = email_lib.utils.formatdate()
    msg.set_content(
        f"Пациент: {patient_name}\n"
        f"Дата: {date_str}\n"
        f"Врач: {doctor}\n\n"
        f"Результаты интерпретации лабораторных анализов во вложении (.docx).\n\n"
        f"Клиника МОЙ МЕД / ЛАЙТМЕД",
        charset="utf-8",
    )
    msg.add_attachment(
        docx_bytes,
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )
    return msg.as_bytes()
